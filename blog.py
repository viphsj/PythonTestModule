#!/usr/bin/env python3
# -*- coding: utf-8 -*-

'a test module'

__author__ = 'Dou Ba'

import os
import shutil
from urllib import request
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_html_soup(url):#获取解编码后的HTML
    html = None
    try:
        response = urllib.request.urlopen(url, timeout = 10)
        html = response.read().decode(encoding = "gb2312", errors='ignore')
    except Exception as e:
        print(e, "please check your network situation")
        return None
    soup = BeautifulSoup(str(html), "lxml")    
    return soup

def page_url(url, page_num):#生成带页面的URL
    if page_num == 1:
        return url
    index = url.rfind(".")
    return url[0 : index] + "_" + str(page_num) + url[index : ]

def get_title_link(url, pattern):#获取新闻的标题和正文链接
    soup = get_html_soup(url)
    news_link = {}

    scroll_list1 = BeautifulSoup(str(soup.find("div", attrs = pattern)), "lxml")
    scroll_list = BeautifulSoup(str(scroll_list1.find_all("h2")), "lxml")
    for link in scroll_list.find_all("a"):    
        if len(link.get_text().strip()) > 0 and link.get("href").find("http") != -1:
            news_link[link.get_text().strip()] = link.get('href')
    return news_link

def get_news_body(url):#抓取新闻主体内容
    first = True
    content_text = []
    page_num = 1
    article_div = ""

    #使用循环处理有分页的新闻
    while first == True or article_div.find("下一页</a>") != -1:
        soup = get_html_soup(page_url(url, page_num))
        if soup == None:
            return None

        #article_div = str(soup.find("div", attrs = {"class": "article"}))
        article_div = str(soup.find("div", attrs = {"class": "news_info"}))
        soup = BeautifulSoup(str(article_div), "lxml")
        for content in soup.find_all("p"):
            if len(content.get_text().strip()) > 0:
                content_text.append(" " + content.get_text().strip())
        page_num += 1
        first = False
    for x in content_text:
        if x == " None":
            return None
    return content_text

def clean_chinese_character(text):
    '''处理特殊的中文符号,将其全部替换为'-' 否则在保存时Windows无法将有的中文符号作为路径'''
    chars = chars = ["/", "\"", "'", "·", "。","？", "！", "，", "、", "；", "：", "‘", "’", "“", "”", "（", "）", "…", "–", "．", "《", "》"];
    new_text = ""
    for i in range(len(text)):
        if text[i] not in chars:
            new_text += text[i]
        else:
            new_text += "_"
    return new_text;

def create_docx(news_type, title, content):
    '''这里使用python-docx库将新闻的内容生成word文件'''
    document = Document()
    paragraph = document.add_paragraph(title)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.bold = True

    for x in content:
        paragraph = document.add_paragraph(x)

    style = paragraph.style
    font = style.font
    font.size = Pt(15)
    font.name = "consolas"

    name = news_type + "-" + clean_chinese_character(title) + ".docx"
    document.save(news_type + "/" + name)

########################################################################

def githrml():
    national = "mydrivers"
    blog_news = "http://blog.mydrivers.com/"
    blog_news_pattern = {"id": "main"}

    #删除旧目录
    print("deleting old dir")
    if os.path.exists(national):
        shutil.rmtree(national)

    #创建新目录
    print("creating dir: ", national)
    os.mkdir(national)

    blog_news_list = get_title_link(blog_news, blog_news_pattern);
    print("\ngetting national news content")
    for x in blog_news_list:
        paras = get_news_body(blog_news_list[x])
        #paras = get_news_body(x)
        if paras != None and len(paras) > 0:
            print("writing:", clean_chinese_character(x), blog_news_list[x])
            create_docx(national, x, paras)

    print("All done, have a nice day")

if __name__ == "__main__":
    githrml()